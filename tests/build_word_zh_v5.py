# -*- coding: utf-8 -*-
"""
build_word_zh_v5.py —— 把 manuscript_zh.md 转成中文 .docx，
                       在每张图 caption 前插入 outputs/figure/ 中的图。
中文字体：正文 宋体、标题 黑体；西文/数字 Arial（Word 自动 fallback）。
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.abspath(os.path.join(BASE, '..'))
MD_PATH = os.path.join(REPO, 'outputs', 'reports', 'manuscript_zh.md')
FIG_DIR = os.path.join(REPO, 'outputs', 'figure')
OUT_PATH = os.path.join(REPO, 'outputs', 'reports',
                        '风电场阵列朝向显著塑造候鸟碰撞暴露（西欧海上与陆上风电场的几何筛查）v6.docx')

FIG_FILES = {
    '1': os.path.join(FIG_DIR, 'fig1.png'),
    '2': os.path.join(FIG_DIR, 'fig2.png'),
    '3': os.path.join(FIG_DIR, 'fig3.png'),
    '4': os.path.join(FIG_DIR, 'fig4.png'),
    '5': os.path.join(FIG_DIR, 'fig5.png'),
    'S1': os.path.join(FIG_DIR, 'figS1.png'),
    'S2': os.path.join(FIG_DIR, 'figS2.png'),
}
FIG_WIDTH = {k: Inches(6.5) for k in FIG_FILES}


def _set_font(run, size=10.5, bold=False, italic=False, cn='宋体', en='Arial'):
    """中英字体分离：run.font.name 用西文字体名，rFonts eastAsia 用中文字体名。"""
    run.font.name = en
    r = run._element
    rpr = r.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), en); rFonts.set(qn('w:hAnsi'), en)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _add_inline_runs(par, text, base_size=10.5, base_bold=False, cn='宋体', en='Arial'):
    tok = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|<sup>[^<]+</sup>)', text)
    for t in tok:
        if not t:
            continue
        if t.startswith('**') and t.endswith('**'):
            r = par.add_run(t[2:-2]); _set_font(r, base_size, bold=True, cn=cn, en=en)
        elif t.startswith('*') and t.endswith('*'):
            r = par.add_run(t[1:-1]); _set_font(r, base_size, bold=base_bold, italic=True, cn=cn, en=en)
        elif t.startswith('`') and t.endswith('`'):
            r = par.add_run(t[1:-1]); _set_font(r, base_size, bold=base_bold, cn=cn, en='Consolas')
        elif t.startswith('<sup>') and t.endswith('</sup>'):
            r = par.add_run(t[5:-6]); _set_font(r, base_size, bold=base_bold, cn=cn, en=en)
            r.font.superscript = True
        else:
            r = par.add_run(t); _set_font(r, base_size, bold=base_bold, cn=cn, en=en)


def _insert_figure(doc, fig_key):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fp = FIG_FILES.get(fig_key)
    if fp and os.path.exists(fp):
        run.add_picture(fp, width=FIG_WIDTH.get(fig_key, Inches(6.5)))
    else:
        run.text = f'[Figure {fig_key} missing]'


def _add_table(doc, tbl_lines):
    rows = []
    for ln in tbl_lines:
        ln = ln.strip()
        if not ln.startswith('|'):
            continue
        cells = [c.strip() for c in ln.strip('|').split('|')]
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncol)
    tbl.style = 'Light Grid Accent 1'
    for ri, r in enumerate(rows):
        for ci in range(ncol):
            txt = r[ci] if ci < len(r) else ''
            cell = tbl.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            _add_inline_runs(p, txt, base_size=9.5, base_bold=(ri == 0))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
    # Normal 样式：中文宋体 / 西文 Arial
    normal = doc.styles['Normal']
    normal.font.name = 'Arial'; normal.font.size = Pt(10.5)
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts'); rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Arial'); rFonts.set(qn('w:hAnsi'), 'Arial')

    lines = open(MD_PATH, encoding='utf-8').read().splitlines()
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        if re.match(r'^-{3,}\s*$', stripped):
            i += 1; continue

        # # 一级标题 -> 论文题目（黑体）
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped[2:].strip())
            _set_font(r, 15, bold=True, cn='黑体')
            i += 1; continue

        # ### 副标题
        if stripped.startswith('### '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped[4:].strip())
            _set_font(r, 12, italic=True, cn='楷体')
            i += 1; continue

        # ## 章节名
        if stripped.startswith('## '):
            p = doc.add_paragraph()
            r = p.add_run(stripped[3:].strip())
            _set_font(r, 13, bold=True, cn='黑体')
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            i += 1; continue

        # 三级 ### 已处理

        # 表格
        if stripped.startswith('|'):
            tbl_lines = []
            while i < n and lines[i].lstrip().startswith('|'):
                tbl_lines.append(lines[i]); i += 1
            _add_table(doc, tbl_lines)
            continue

        if stripped == '':
            i += 1; continue

        # 图 caption：以 **Figure N | 或 **Supplementary Fig. SN | 开头
        m = re.match(r'^\*\*(Figure|Supplementary Fig\.)\s+(S?\d+)\s*\|', stripped)
        if m:
            fig_key = m.group(2)
            _insert_figure(doc, fig_key)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(12)
            _add_inline_runs(p, stripped, base_size=9.5)
            i += 1; continue

        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Pt(21)  # 首行缩进 2 字符
        _add_inline_runs(p, stripped, base_size=10.5)
        i += 1

    doc.save(OUT_PATH)
    print('wrote:', OUT_PATH)


if __name__ == '__main__':
    build()
