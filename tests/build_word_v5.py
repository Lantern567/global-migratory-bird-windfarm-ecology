# -*- coding: utf-8 -*-
"""
build_word_v5.py —— 把 manuscript_en.md 转成 .docx，并在每张主图/附图 caption 后
                     插入 outputs/figure/fig1..fig5.png / figS1.png / figS2.png。
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

BASE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.abspath(os.path.join(BASE, '..'))
MD_PATH = os.path.join(REPO, 'outputs', 'reports', 'manuscript_en.md')
FIG_DIR = os.path.join(REPO, 'outputs', 'figure')
OUT_PATH = os.path.join(REPO, 'outputs', 'reports',
                        'Wind farm array orientation strongly shapes migratory bird collision exposure (v6).docx')

# 图号 -> 文件（主图 fig1..fig5，附图 figS1/S2）
FIG_FILES = {
    '1': os.path.join(FIG_DIR, 'fig1.png'),
    '2': os.path.join(FIG_DIR, 'fig2.png'),
    '3': os.path.join(FIG_DIR, 'fig3.png'),
    '4': os.path.join(FIG_DIR, 'fig4.png'),
    '5': os.path.join(FIG_DIR, 'fig5.png'),
    'S1': os.path.join(FIG_DIR, 'figS1.png'),
    'S2': os.path.join(FIG_DIR, 'figS2.png'),
}
# Word 中图宽（Nature Energy 双栏 = 6.5"；单栏附图也用双栏方便阅读）
FIG_WIDTH = {
    '1': Inches(6.5), '2': Inches(6.5), '3': Inches(6.5), '4': Inches(6.5), '5': Inches(6.5),
    'S1': Inches(6.5), 'S2': Inches(6.5),
}


def _set_font(run, size=10.5, bold=False, italic=False, color=None):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _add_inline_runs(par, text, base_size=10.5, base_bold=False):
    """支持行内 **bold** / *italic* / `code` / <sup>n</sup> 的简易分词。"""
    tok = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|<sup>[^<]+</sup>)', text)
    for t in tok:
        if not t:
            continue
        if t.startswith('**') and t.endswith('**'):
            r = par.add_run(t[2:-2]); _set_font(r, base_size, bold=True)
        elif t.startswith('*') and t.endswith('*'):
            r = par.add_run(t[1:-1]); _set_font(r, base_size, bold=base_bold, italic=True)
        elif t.startswith('`') and t.endswith('`'):
            r = par.add_run(t[1:-1]); _set_font(r, base_size, bold=base_bold)
            r.font.name = 'Consolas'
        elif t.startswith('<sup>') and t.endswith('</sup>'):
            r = par.add_run(t[5:-6]); _set_font(r, base_size, bold=base_bold)
            r.font.superscript = True
        else:
            r = par.add_run(t); _set_font(r, base_size, bold=base_bold)


def _insert_figure(doc, fig_key):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fp = FIG_FILES.get(fig_key)
    if fp and os.path.exists(fp):
        run.add_picture(fp, width=FIG_WIDTH.get(fig_key, Inches(6.5)))
    else:
        run.text = f'[Figure {fig_key} missing]'


def build():
    doc = Document()
    # 页面：Nature Energy 主稿常规 A4，边距 2 cm
    for s in doc.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10.5)

    lines = open(MD_PATH, encoding='utf-8').read().splitlines()
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        # 忽略仅由 --- 组成的水平线
        if re.match(r'^-{3,}\s*$', stripped):
            i += 1; continue

        # 一级标题 # -> 论文题目
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped[2:].strip())
            _set_font(r, 14, bold=True)
            i += 1; continue

        # 三级标题 ### -> 副标题
        if stripped.startswith('### '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped[4:].strip())
            _set_font(r, 12, italic=True)
            i += 1; continue

        # 二级标题 ## -> 章节名
        if stripped.startswith('## '):
            p = doc.add_paragraph()
            r = p.add_run(stripped[3:].strip())
            _set_font(r, 12, bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            i += 1; continue

        # 表格：连续以 | 开头的行
        if stripped.startswith('|'):
            tbl_lines = []
            while i < n and lines[i].lstrip().startswith('|'):
                tbl_lines.append(lines[i]); i += 1
            _add_table(doc, tbl_lines)
            continue

        # 空行
        if stripped == '':
            i += 1; continue

        # 普通段落 / 图表 caption / 参考文献编号项
        # 检查是否是图 caption：以 **Figure N | 或 **Supplementary Fig. SN | 开头
        m = re.match(r'^\*\*(Figure|Supplementary Fig\.)\s+(S?\d+)\s*\|', stripped)
        if m:
            # 先插图，再写 caption（Nature 期刊惯例 caption 位于图下方）
            fig_key = m.group(2)
            _insert_figure(doc, fig_key)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(12)
            _add_inline_runs(p, stripped, base_size=9.5)
            i += 1; continue

        # 参考文献编号 "1. ..." 或列表项 "- ..." 保持段落
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _add_inline_runs(p, stripped, base_size=10.5)
        i += 1

    doc.save(OUT_PATH)
    print('wrote:', OUT_PATH)


def _add_table(doc, tbl_lines):
    # 去掉分隔行 |---|---|
    rows = []
    for ln in tbl_lines:
        ln = ln.strip()
        if not ln.startswith('|'):
            continue
        cells = [c.strip() for c in ln.strip('|').split('|')]
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncol)
    tbl.style = 'Light Grid Accent 1'
    for ri, r in enumerate(rows):
        for ci in range(ncol):
            txt = r[ci] if ci < len(r) else ''
            cell = tbl.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            _add_inline_runs(p, txt, base_size=9.5, base_bold=(ri == 0))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


if __name__ == '__main__':
    build()
