# -*- coding: utf-8 -*-
"""
build_word_v7.py —— 将 manuscript_en.md 编译为 Nature Energy 版 Word（v7），
                    图片来自 figures_v6/（新的多子图重设计），并动态替换过时的
                    figure caption 段落以匹配 v6 的 panel 结构。

改动与 v5 差异：
  * 图源目录：figures_v6/*.png（而非 outputs/figure）
  * 图 caption 采用 v7 panel 结构（sin² surface / Pareto frontier / ridgeline
    / hex-bin / case-study / saturation heatmap 等）
  * 移除 Table 2（并入 Fig 5i heatmap）
  * 输出至 outputs/reports/…v7.docx
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

BASE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.abspath(os.path.join(BASE, '..'))
MD_PATH = os.path.join(REPO, 'outputs', 'reports', 'manuscript_en.md')
FIG_DIR = os.path.join(REPO, 'figures_v6')
OUT_PATH = os.path.join(REPO, 'outputs', 'reports',
                        'Wind farm array orientation strongly shapes migratory bird collision exposure (v7).docx')

FIG_FILES = {
    # Fig 1 已于 2026-08-28 重做为 figures_v7 的「全幅地图 + 环绕机理链」版，
    # 不再来自 figures_v6；其余图仍取 v6。
    '1': os.path.join(REPO, 'figures_v7', 'fig1_map_flow.png'),
    '2': os.path.join(REPO, 'figures_v7', 'fig2_sensitivity.png'),
    '3': os.path.join(REPO, 'figures_v7', 'fig3_misalignment.png'),
    '4': os.path.join(REPO, 'figures_v7', 'fig4_capture.png'),
    '5': os.path.join(REPO, 'figures_v7', 'fig5_tradeoff.png'),
    'S1': os.path.join(FIG_DIR, 'figS1_threat.png'),
    'S2': os.path.join(FIG_DIR, 'figS2_universality.png'),
}
FIG_WIDTH = {k: Inches(6.5) for k in FIG_FILES}

# 新版 caption（英文，对应 v6 panel 结构）
NEW_CAPTIONS_EN = {
    '1': (
        '**Figure 1 | Study region and the orientation mechanism that governs migratory collision exposure.** **a**, Western European study region, shown full-bleed: arrows give the spring migration direction field (2,025 grid cells, one arrow drawn per second cell; colour = directional concentration, 0.84–1.00), overlaid with 4,191 onshore wind farms, 29 offshore farms with VPTS weather-radar bird directions, 26 offshore farms on the Bauer grid, and 6 VPTS radar stations. **b**–**f**, the mechanism chain, arranged around the map over areas that contain no study data, so that no wind farm or migration vector is occluded. **b**, The migration field sets the reference direction φ: the circular mean spring bearing is 52° and the median directional concentration is 0.89. **c**, Array orientation θ relative to φ sets geometric exposure — birds crossing parallel to the turbine rows (θ ∥ φ) pass along the inter-row corridors and are minimally exposed, whereas birds crossing perpendicular (θ ⊥ φ) traverse the rotor-swept plane once per row — giving E ∝ sin²(θ − φ). **d**, The energy optimum and the ecological optimum do not coincide: the median misalignment Δ between the AEP-optimal and exposure-minimal orientations is 48° onshore, 40° for VPTS farms and 45° for Bauer farms (bars, top to bottom, span the interquartile range: 29–60°, 20–50° and 30–67.5°). **e**, Most of the avoidable exposure is recovered by a small rotation: group medians of the fraction of avoidable exposure removed as the array is rotated by Δθ away from its AEP optimum. Half of the maximum gain is reached at Δθ₅₀ = 17° onshore (14° VPTS, 17.5° Bauer), and a rotation of ≤20° (dashed line) already removes 55% of the avoidable exposure onshore. **f**, That rotation is nearly free: median exposure cut attainable within an AEP budget of 0.5, 1, 2 and 5%. At the 1% budget (box) the median cut is 97% onshore, 84% for VPTS and 48% for Bauer farms, while the AEP loss actually incurred is only 0.54%, 0.28% and 0.53%, respectively.'
    ),
    '2': (
        "**Figure 2 | Migration keeps a stable axis, and that axis makes collision exposure highly sensitive to array orientation.** Panels are grouped by the claim they support. **a–c, migration keeps a stable axis. a**, Flight-direction rose over the 2,025 Bauer grid cells; arrows mark the circular median (spring 52°, autumn 235°). **b**, Directional concentration (median 0.89 in spring, 0.76 in autumn). **c**, The same directions folded onto the 180°-periodic exposure axis: the two seasons reverse as vectors but their exposure axes differ by only 3.4°, so a single orientation can suppress both seasons at once. **d–f, that axis makes exposure sensitive to orientation. d**, Family of exposure curves, each farm aligned on its own minimum (group median and 10–90% band), against the analytic sin² geometry. **e**, Per-farm sensitivity amplitude — the share of a farm's exposure range that rotation alone can span (median 99.9% onshore, 97.6% VPTS, 93.4% Bauer; every farm exceeds 80%); half-violin, interquartile bar and individual farms. **f**, Mechanism: the more concentrated the local flight direction, the wider the reachable range (onshore, binned median ± IQR). **g,h, the sensitivity is universal. g**, Sensitivity across space (onshore hexagon medians, the 55 offshore farms overlaid). **h**, The three independent data sources agree within a 93–100% band; circles are medians, bars the interquartile range, whiskers the 5–95% range."
    ),
    '3': (
        '**Figure 3 | The energy-optimal and the exposure-optimal orientation are systematically misaligned.** **a–c, the two optima land apart. a**, The AEP-optimal orientation is set by the wind resource and spreads across every angle, whereas the exposure-optimal orientation is pinned to the migration axis (all 4,246 farms, back-to-back histograms). **b**, Per-farm misalignment Δ (median 48° onshore, 40° VPTS, 45° Bauer). **c**, Joint distribution of the two optima; the dashed line marks perfect alignment and the vertical axis is clipped to the range the data occupy. **d–f, the exposure penalty is an order of magnitude. d**, Ratio of the exposure at the AEP optimum to the reachable minimum (median 603× onshore, 19× VPTS, 10× Bauer); the axis is clipped at 10⁵ because E_min approaches zero for purely geometric farms. **e**, A representative onshore farm with Δ near the median — exposure (vermillion) is minimal where AEP (grey) is not. **f**, The ratio grows with misalignment (onshore binned median ± IQR, offshore farms overlaid). **g–i, nearly all of that excess is removable. g**, Avoidable exposure share (median 99.8% / 94.6% / 90.0%). **h**, Share of farms at or above a given avoidable share. **i**, Avoidable share across space.'
    ),
    '4': (
        '**Figure 4 | A limited rotation captures most of the avoidable exposure.** **a–c, the benefit is front-loaded. a**, Capture curves (median and interquartile band) as the array is rotated away from its AEP optimum; shading marks the first 20°. **b**, Median share recovered at 5°, 10°, 20° and 30° of rotation. **c**, Marginal return per extra degree, which falls away beyond about 20°. **d–f, half the gain needs about a third of the turn. **d**, Group medians of Δθ50, Δθ80 and the full misalignment. **e**, Per-farm Δθ50 against full misalignment, with the 1:1 and one-third references. **f**, Δθ50 across space (onshore hexagon medians). **g,h, the rule holds farm by farm. g**, Share of farms recovering 50% or 80% of the avoidable exposure within 20° or 30°. **h**, Cumulative distribution of Δθ50; dots mark the share reached at 20° and 30°.'
    ),
    '5': (
        '**Figure 5 | The ecological gain costs almost no energy.** **a–c, a 1% allowance buys most of the cut. a**, Energy–exposure trade-off for every farm and budget; large markers are per-group medians at the 1% budget. **b**, Mean energy paid against mean exposure bought at the 1% budget, on a logarithmic axis because the two differ by two orders of magnitude. **c**, Distribution of the exposure reduction achieved at the 1% budget (median 97% onshore, 84% VPTS, 48% Bauer). **d,e, the asymmetry is geometric. **d**, Around its own optimum AEP is flat while exposure is steep (onshore, median ± IQR over 500 farms), so a small energy concession buys a large exposure cut. **e**, Median exposure reduction against the AEP budget. **f–h, the return dies out past about 1%. f**, Share of farms that gain nothing when the budget is raised from 2% to 5%. **g**, Policy scenarios: share of farms meeting an ecological floor (≥50%, ≥80%, ≥90% exposure reduction; line style) under each AEP budget ceiling (colour = group). The VPTS ≥90% line is flat from the 1% budget onward — beyond that point a larger energy concession buys no additional farms. **h**, Exposure reduction at the 1% budget across space.'
    ),
}

# 中文版对应 caption（关键描述与英文一致）
NEW_CAPTIONS_ZH = {
    '1': (
        '**Figure 1 | 研究区与决定候鸟碰撞暴露的朝向机制。** **a**，西欧研究区（铺满全幅）：箭矢为春季迁徙方向场（2,025 个格网单元，隔一格绘一支；颜色为方向集中度，0.84–1.00），叠加 4,191 个陆上风电场、29 个具 VPTS 天气雷达鸟类方向的海上风电场、26 个 Bauer 格网海上风电场与 6 个 VPTS 雷达站。**b**–**f**，机理链，沿地图外圈布置于不含任何研究数据的区域，因此未遮挡任何风电场或迁徙矢量。**b**，迁徙场给定参考方向 φ：春季方向的圆均值为 52°，方向集中度中位数为 0.89。**c**，阵列朝向 θ 相对 φ 决定几何暴露：与机组行方向平行穿越（θ ∥ φ）的鸟可沿行间通道通过，暴露极小；垂直穿越（θ ⊥ φ）则须每行穿越一次旋翼扫掠面，故 E ∝ sin²(θ − φ)。**d**，能源最优与生态最优并不重合：AEP 最优朝向与暴露最小朝向的错位中位数为陆上 48°、VPTS 40°、Bauer 45°（横条自上而下为三组的四分位距，分别为 29–60°、20–50°、30–67.5°）。**e**，有限旋转即可收回大部分可削减暴露：三组中位曲线给出阵列自 AEP 最优朝向旋转 Δθ 后已削减的可避免暴露份额。达到最大收益一半所需的 Δθ₅₀ 为陆上 17°（VPTS 14°、Bauer 17.5°）；旋转 ≤20°（虚线）已可削减陆上 55% 的可避免暴露。**f**，这一旋转的能源代价极低：0.5%、1%、2%、5% 四档AEP 预算下可实现的暴露削减中位数。在 1% 预算（方框）下，陆上、VPTS、Bauer 的削减中位数分别为 97%、84%、48%，而实际发生的 AEP 损失仅为 0.54%、0.28%、0.53%。'
    ),
    '2': (
        '**Figure 2 | 迁徙方向具有稳定主轴，该主轴使碰撞暴露对阵列朝向高度敏感。** 子图按所支撑的子论点分带。**a–c，迁徙有稳定主轴。a**，2,025 个 Bauer 格点的飞行方向玫瑰图，箭头为圆中位（春 52°、秋 235°）。**b**，方向集中度分布（春季中位 0.89、秋季 0.76）。**c**，同一批方向折叠到以 180° 为周期的暴露轴上：两季作为向量近乎相反，但作为暴露轴仅相差 3.4°，因此单一朝向即可同时压低两季暴露。**d–f，该主轴使暴露对朝向高度敏感。d**，暴露曲线族，各场以自身最小值对齐（分组中位与 10–90% 带），并与解析的 sin² 几何对照。**e**，逐场方向敏感性幅度 —— 仅靠旋转即可覆盖的暴露区间份额（中位陆上 99.9%、VPTS 97.6%、Bauer 93.4%；所有场址均超过 80%）；半小提琴、四分位条与逐场散点。**f**，机制：当地飞行方向越集中，可达区间越宽（陆上，分箱中位 ± 四分位距）。**g,h，该敏感性具有普遍性。g**，敏感性的空间分布（陆上六边形中位，叠加 55 个海上场）。**h**，三套独立数据源一致落在 93–100% 区间；圆点为中位、粗条为四分位距、细线为 5–95% 区间。'
    ),
    '3': (
        '**Figure 3 | 能源最优朝向与暴露最优朝向系统性错位。** **a–c，两个最优彼此错开。a**，AEP 最优朝向由风资源决定、散布于全部角域，而暴露最优朝向被迁徙轴钉住（全部 4,246 场，背靠背直方图）。**b**，逐场错位角 Δ（中位陆上 48°、VPTS 40°、Bauer 45°）。**c**，两个最优的联合分布，虚线为完全对齐；纵轴按数据实际占据的范围裁剪。**d–f，能源最优处的暴露高出一个数量级。**d**，AEP 最优处暴露与可达最小暴露之比（中位陆上 603×、VPTS 19×、Bauer 10×）；纯几何场址的E_min 趋近于零，故坐标轴在 10⁵ 处裁剪。**e**，错位角接近中位的代表性陆上场址 —— 暴露（朱红）的极小值并不落在 AEP（灰）的极大值处。**f**，该比值随错位角上升（陆上分箱中位 ± 四分位距，叠加海上场）。**g–i，这部分超额暴露绝大多数可消除。g**，可削减暴露比例（中位 99.8% / 94.6% / 90.0%）。**h**，达到或超过给定可削减比例的场址份额。**i**，可削减比例的空间分布。'
    ),
    '4': (
        '**Figure 4 | 有限的朝向调整即可捕获大部分可削减暴露。** **a–c，收益集中在前段。a**，自 AEP 最优起旋转的捕获曲线（中位与四分位带），阴影为最初 20°。**b**，旋转 5°、10°、20°、30° 时回收的中位份额。**c**，每多转一度的边际收益，约 20° 之后迅速衰减。**d–f，半收益仅需约三分之一的转角。d**，各组 Δθ50、Δθ80 与全错位角的中位数。**e**，逐场 Δθ50 对全错位角，并给出 1:1与三分之一参考线。**f**，Δθ50 的空间分布（陆上六边形中位）。**g,h，该规律在场址层面成立。**g**，在 20° 或 30° 内回收 50% 或 80% 可削减暴露的场址比例。**h**，Δθ50 的累积分布，圆点标出20° 与 30° 处的达成比例。'
    ),
    '5': (
        '**Figure 5 | 生态收益的能源代价几乎为零。** **a–c，1% 预算即可买到大部分削减。a**，全部场址 × 全部预算的能源–暴露权衡，大标记为各组在 1% 预算下的中位。**b**，1% 预算下平均付出的能源与平均买到的暴露削减；两者相差两个数量级，故用对数轴。**c**，1% 预算下实现的暴露削减分布（中位陆上 97%、VPTS 84%、Bauer 48%）。**d,e，不对称源于几何。d**，在各自最优点附近，AEP 是平的而暴露是陡的（陆上，500 场的中位 ± 四分位距），因此极小的能源让步可换得大幅暴露削减。**e**，暴露削减中位随 AEP 预算的变化。**f–h，约 1% 之后边际回报消失。f**，预算由 2% 提高到5% 时毫无新增收益的场址比例。**g**，政策情境：在各档 AEP 预算上限（横轴）下达到给定生态底线（≥50%、≥80%、≥90% 暴露削减；线型区分）的场址占比，颜色区分三组。VPTS 的 ≥90% 曲线自 1% 预算起走平 —— 越过该点后，更大的能源让步不再换来更多达标场址。**h**，1% 预算下暴露削减的空间分布。'
    ),
}

# Table 2 现由 Fig 5i 替代 —— 从 md 里删除相关行
TABLE2_HEADER_RE = re.compile(r'^\*\*Table 2 \|', re.IGNORECASE)


def _set_font(run, size=10.5, bold=False, italic=False, color=None):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _add_inline_runs(par, text, base_size=10.5, base_bold=False):
    tok = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|<sup>[^<]+</sup>)', text)
    for t in tok:
        if not t:
            continue
        if t.startswith('**') and t.endswith('**'):
            r = par.add_run(t[2:-2]); _set_font(r, base_size, bold=True)
        elif t.startswith('*') and t.endswith('*'):
            r = par.add_run(t[1:-1]); _set_font(r, base_size, bold=base_bold, italic=True)
        elif t.startswith('`') and t.endswith('`'):
            r = par.add_run(t[1:-1]); _set_font(r, base_size, bold=base_bold)
            r.font.name = 'Consolas'
        elif t.startswith('<sup>') and t.endswith('</sup>'):
            r = par.add_run(t[5:-6]); _set_font(r, base_size, bold=base_bold)
            r.font.superscript = True
        else:
            r = par.add_run(t); _set_font(r, base_size, bold=base_bold)


def _insert_figure(doc, fig_key):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fp = FIG_FILES.get(fig_key)
    if fp and os.path.exists(fp):
        run.add_picture(fp, width=FIG_WIDTH.get(fig_key, Inches(6.5)))
    else:
        run.text = f'[Figure {fig_key} missing]'


def _add_table(doc, tbl_lines):
    rows = []
    for ln in tbl_lines:
        ln = ln.strip()
        if not ln.startswith('|'):
            continue
        cells = [c.strip() for c in ln.strip('|').split('|')]
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncol)
    tbl.style = 'Light Grid Accent 1'
    for ri, r in enumerate(rows):
        for ci in range(ncol):
            txt = r[ci] if ci < len(r) else ''
            cell = tbl.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            _add_inline_runs(p, txt, base_size=9.5, base_bold=(ri == 0))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _remove_table2_block(lines):
    """去掉 Table 2 段落 + 后续表格（改为 Fig 5i heatmap 替代）。"""
    out = []
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        if TABLE2_HEADER_RE.match(line.lstrip()):
            # 跳过：Table 2 caption 段
            i += 1
            # 跳过：空行
            while i < n and lines[i].strip() == '':
                i += 1
            # 跳过：表格（连续 | 开头的行）
            while i < n and lines[i].lstrip().startswith('|'):
                i += 1
            # 跳过：紧随其后的空行
            while i < n and lines[i].strip() == '':
                i += 1
            continue
        out.append(line)
        i += 1
    return out


def build(md_path, out_path, captions):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10.5)

    lines = open(md_path, encoding='utf-8').read().splitlines()
    lines = _remove_table2_block(lines)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        if re.match(r'^-{3,}\s*$', stripped):
            i += 1; continue

        if stripped.startswith('# '):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped[2:].strip()); _set_font(r, 14, bold=True)
            i += 1; continue
        if stripped.startswith('### '):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped[4:].strip()); _set_font(r, 12, italic=True)
            i += 1; continue
        if stripped.startswith('## '):
            p = doc.add_paragraph()
            r = p.add_run(stripped[3:].strip()); _set_font(r, 12, bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            i += 1; continue

        if stripped.startswith('|'):
            tbl_lines = []
            while i < n and lines[i].lstrip().startswith('|'):
                tbl_lines.append(lines[i]); i += 1
            _add_table(doc, tbl_lines)
            continue

        if stripped == '':
            i += 1; continue

        # 图/附图 caption：替换主图 caption 为 v7 版
        m = re.match(r'^\*\*(Figure|Supplementary Fig\.)\s+(S?\d+)\s*\|', stripped)
        if m:
            fig_key = m.group(2)
            _insert_figure(doc, fig_key)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(12)
            cap_text = captions.get(fig_key, stripped)
            _add_inline_runs(p, cap_text, base_size=9.5)
            i += 1; continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _add_inline_runs(p, stripped, base_size=10.5)
        i += 1

    doc.save(out_path)
    print('wrote:', out_path)


if __name__ == '__main__':
    # 英文
    build(MD_PATH, OUT_PATH, NEW_CAPTIONS_EN)
    # 中文
    zh_md = os.path.join(REPO, 'outputs', 'reports', 'manuscript_zh.md')
    zh_out = os.path.join(REPO, 'outputs', 'reports',
                          '风电场阵列朝向显著塑造候鸟碰撞暴露（西欧海上与陆上风电场的几何筛查）v7.docx')
    if os.path.exists(zh_md):
        build(zh_md, zh_out, NEW_CAPTIONS_ZH)
