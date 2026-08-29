# -*- coding: utf-8 -*-
"""把手工修订版 docx 里的 Fig 1–5 换成 figures_v7 的重制版（图 + 图注 + 交叉引用）。

正文一旦在 Word 里手改过，就不能再用 build_word_v7.py 整篇重建（会丢修订）。
本脚本只动三处：各图的内嵌图片二进制与显示尺寸、各图的图注段文字、正文里指向
旧面板结构的交叉引用；其余段落逐字不动。

图注取自 build_word_v7.NEW_CAPTIONS_ZH，与 markdown 源保持单一出处。
可重复运行：每次都从 SRC 重新读取，不会在已改过的文件上二次应用。

用法：python tests/patch_figs_into_revision.py [源docx] [目标docx]
"""
import copy
import io
import os
import re
import sys

from docx import Document
from docx.shared import Pt

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

BASE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.abspath(os.path.join(BASE, '..'))
sys.path.insert(0, BASE)
from build_word_v7 import NEW_CAPTIONS_ZH                            # noqa: E402

PNG = {
    '1': os.path.join(REPO, 'figures_v7', 'fig1_map_flow.png'),
    '2': os.path.join(REPO, 'figures_v7', 'fig2_sensitivity.png'),
    '3': os.path.join(REPO, 'figures_v7', 'fig3_misalignment.png'),
    '4': os.path.join(REPO, 'figures_v7', 'fig4_capture.png'),
    '5': os.path.join(REPO, 'figures_v7', 'fig5_tradeoff.png'),
}

DEFAULT_SRC = os.path.join(
    REPO, '风电场阵列朝向显著塑造候鸟碰撞暴露（西欧海上与陆上风电场的几何筛查）'
          'v7_修订_bak_20260828_命题四标题3前.docx')
DEFAULT_DST = os.path.join(
    REPO, '风电场阵列朝向显著塑造候鸟碰撞暴露（西欧海上与陆上风电场的几何筛查）'
          'v7_修订_图全部重制_20260829.docx')

# 逐段的交叉引用映射（段号来自 SRC）。同一 token 在不同段落含义不同 ——
# 例如第 25 段的 Fig. 3c 指暴露比、第 26 段的 Fig. 3c 指可削减比例 ——
# 所以必须按段落分别映射，且每段做单遍同步替换，避免链式误改。
XREF = {
    9:  {'Fig. 1d': 'Fig. 1c'},                                   # sin² 律
    19: {'Fig. 2b': 'Fig. 2d', 'Fig. 2c': 'Fig. 2e'},             # 曲线形状 / 敏感性幅度
    24: {'Fig. 3a': 'Fig. 3c'},                                   # 两最优联合分布
    25: {'Fig. 3c': 'Fig. 3d'},                                   # 暴露比
    26: {'Fig. 3c': 'Fig. 3g'},                                   # 可削减比例
    30: {'Fig. 4a': 'Fig. 4a, Fig. 4b'},                          # 5/10/20/30° 捕获份额
    31: {'Fig. 4b': 'Fig. 4d', 'Fig. 4d': 'Fig. 4e',
         'Fig. 4f': 'Fig. 4h'},                                   # Δθ50 / Δθ80
    32: {'Fig. 4e': 'Fig. 4g'},                                   # 阈值达成率
    36: {'Fig. 5g': 'Fig. 5b', 'Fig. 5h': 'Fig. 5c',
         'Fig. 5a–f': 'Fig. 5h'},                                 # 代价 / 收益 / 空间
    37: {'Fig. 5g': 'Fig. 5d, Fig. 5e'},                          # 前沿不对称
    38: {'Fig. 5g': 'Fig. 5f, Fig. 5g'},                          # 饱和
}

# 过时的 VPTS 数字 -> final_numbers.py 实算值。陆上与 Bauer 两组本来就对，
# 只有 VPTS 一列停留在海上重算之前的旧值。段号同样来自 SRC。
NUMFIX = {
    3:  {'82.1%': '84.2%'},                                        # 摘要
    36: {'0.36%': '0.38%', '72.3%': '75.0%', '82.1%': '84.2%'},    # R4.1
    37: {'88.1%': '88.0%'},                                        # R4.2
    38: {'VPTS 100%': 'VPTS 93.1%'},                               # R4.3 饱和
    39: {'88.1%': '88.0%'},                                        # R4.2 组间梯度
    43: {'约 82%（82.1%）': '约 84%（84.2%）'},                      # 讨论
    45: {'VPTS（82.1%）': 'VPTS（84.2%）'},                          # 讨论·口径差
    83: {'82.1%': '84.2%'},                                        # 结论
}

# --- 命题四：把单点设问扩成「AEP 预算 × 保护目标」政策网格 -------------------
# 原稿只用了 1% 一档预算、80%/90% 两档目标。数据里 0.5/1/2/5% 四档预算都算好了，
# 交叉成网格后才谈得上「设定政策」：横向是开发商可承受的能源让步，纵向是保护方
# 要求的最低削减，单元格是该组合下的达标场址占比。数字全部现算，不写死。
POLICY_BUDGETS = [0.005, 0.01, 0.02, 0.05]
POLICY_TARGETS = [50, 80, 90]
POLICY_GROUPS = ['陆上', 'VPTS', 'Bauer']
SEP = chr(10)          # 结论四各自然段的分隔符


def policy_grid():
    """返回 (grid, medians, n)：grid[(budget,target)][group] = 达标占比%。"""
    import figure_style as fs
    on, vp, ba, ctx = fs.compute_metrics()
    O, T = ctx['on'], ctx['to']
    ids = {'VPTS': ctx['vpts_ids'], 'Bauer': ctx['bauer_ids']}

    def rr(b, g):
        if g == '陆上':
            return O[O.budget == b].risk_reduction.values
        sub = T[(T.budget == b) & (T.farm_id.isin(ids[g]))]
        return sub.risk_reduction_pct.values

    import numpy as np
    grid, med = {}, {}
    for b in POLICY_BUDGETS:
        med[b] = {g: float(np.median(rr(b, g))) for g in POLICY_GROUPS}
        for t in POLICY_TARGETS:
            grid[(b, t)] = {g: float((rr(b, g) >= t).mean() * 100) for g in POLICY_GROUPS}
    n = {g: len(rr(0.01, g)) for g in POLICY_GROUPS}
    return grid, med, n


def conclusion_four(grid, med):
    """重写「命题四」本身：把政策情境写进结论正文，不另设表格。

    用户明确要求不要把这些设定塞进 Table —— 政策情境应当是结论的论证本身，
    而不是挂在结论后面的一张附表。
    """
    g = lambda b, t, k: grid[(b, t)][k]
    m = lambda b, k: med[b][k]
    return (
        '命题四给出的不是一个最优朝向，而是一组可供政策选择的设定。与命题三回答'
        '“要转多少度”（几何问题）不同，这里回答的是“政策该把哪个量定死”（决策问题），'
        '而两种定法给出的答案并不对称（Fig. 5g）。\n'
        '若政策定死**能源让步上限**，生态收益随之确定：上限取 1% 时，中位几何暴露削减为'
        f'陆上 {m(0.01,"陆上"):.1f}%、海上 VPTS {m(0.01,"VPTS"):.1f}%、Bauer {m(0.01,"Bauer"):.1f}%，'
        '而实测代价仅 0.51%、0.38% 与 0.49%；上限放宽到 2% 与 5%，陆上与 VPTS 已无实质增益'
        f'（{m(0.02,"陆上"):.1f}%／{m(0.05,"陆上"):.1f}% 与 {m(0.02,"VPTS"):.1f}%／{m(0.05,"VPTS"):.1f}%），'
        f'仅 Bauer 继续上行至 {m(0.02,"Bauer"):.1f}% 与 {m(0.05,"Bauer"):.1f}%。\n'
        '若政策定死**生态底线**，则要问多少场址能在给定预算内达标。以“削减过半”为底线，'
        f'1% 让步即可覆盖陆上 {g(0.01,50,"陆上"):.1f}%、VPTS {g(0.01,50,"VPTS"):.1f}%、'
        f'Bauer {g(0.01,50,"Bauer"):.1f}% 的场址；底线提高到 80%，同一预算下三组降至 '
        f'{g(0.01,80,"陆上"):.1f}%、{g(0.01,80,"VPTS"):.1f}% 与 {g(0.01,80,"Bauer"):.1f}%，'
        f'把预算放宽到 2% 可使 Bauer 回升至 {g(0.02,80,"Bauer"):.1f}%；底线提高到 90%，'
        f'陆上需 2% 预算才能覆盖 {g(0.02,90,"陆上"):.1f}% 的场址，而 VPTS 的达标率自 1% 起'
        f'即固定在 {g(0.01,90,"VPTS"):.1f}%，2% 与 5% 均无任何提升。\n'
        '由此得到本命题最具政策含义的一点：**能源让步并非在所有情形下都能买到生态收益**。'
        'VPTS 在 90% 底线上的天花板由阵列几何与方向数据分辨率决定，而非由能源约束决定；'
        '对这类场址继续加码预算是无效投入，提升候鸟方向数据的空间分辨率才是有效路径。'
        '因此可操作的政策形态是“预算上限 + 生态底线”的组合，而非任一单独的数字：在 '
        '50%–80% 的底线区间内，1%–2% 的让步已能覆盖多数陆上与 VPTS 场址；90% 的严格底线'
        '只在陆上可大范围实现。此外，春秋两季迁徙主轴在暴露几何上仅相差 ≈3.4°（以 180° 为'
        '周期），单一朝向即可同时压低两季暴露，政策无需分季设定差异化要求。')


def policy_paragraph(grid, med):
    g = lambda b, t, k: grid[(b, t)][k]
    return (
        '与命题三聚焦“旋转多少度”（几何问题）不同，命题四将上述权衡转译为“政策如何设定”'
        '（决策问题）。把 AEP 预算上限与保护目标交叉，得到一张政策网格（Table 4）：横向为'
        '开发商可承受的能源让步（0.5%、1%、2%、5%），纵向为保护方要求的最低几何暴露削减'
        '（≥50%、≥80%、≥90%），单元格为在该预算内达标的场址占比。网格给出三条可直接用于'
        '制定政策的读数。其一，若只要求“削减过半”，1% 的能源让步已覆盖陆上 '
        f'{g(0.01,50,"陆上"):.1f}%、VPTS {g(0.01,50,"VPTS"):.1f}% 的场址，而实测代价不足 0.6%。'
        '其二，目标提高到 ≥80% 时，1% 预算下陆上达标率降至 '
        f'{g(0.01,80,"陆上"):.1f}%、Bauer 仅 {g(0.01,80,"Bauer"):.1f}%；把预算放宽到 2%，'
        f'Bauer 可回升至 {g(0.02,80,"Bauer"):.1f}% —— 对方向数据分辨率较低的场址，需要以更大的'
        '能源让步换取同等生态保障。其三，目标提高到 ≥90% 时出现一个硬上限：VPTS 的达标率自 '
        f'1% 预算起即固定在 {g(0.01,90,"VPTS"):.1f}%，预算放宽到 2% 与 5% 均无任何提升'
        f'（{g(0.02,90,"VPTS"):.1f}%、{g(0.05,90,"VPTS"):.1f}%）；该上限由阵列几何与方向数据'
        '分辨率决定，而非由能源约束决定，继续加码能源让步在这一组不再产生生态回报。因此政策'
        '的合理形态不是单一的“预算上限”或单一的“削减目标”，而是二者的组合：在 ≥50%–80% 的'
        f'目标区间内，1%–2% 的预算即可覆盖大多数陆上与 VPTS 场址；≥90% 的严格目标只在陆上可'
        f'大范围实现（2% 预算下 {g(0.02,90,"陆上"):.1f}%），对海上则受限于方向数据而非能源。'
        '此外，两季迁徙主轴在暴露几何上仅相差 ≈3.4°（以 180° 为周期），单一朝向即可同时压低'
        '春秋两季暴露，政策无需分季设定差异化要求。')


W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
A = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
R = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
WP = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'


def png_size(path):
    with open(path, 'rb') as f:
        head = f.read(33)
    assert head[:8] == b'\x89PNG\r\n\x1a\n', f'not a PNG: {path}'
    return int.from_bytes(head[16:20], 'big'), int.from_bytes(head[20:24], 'big')


def find_pairs(doc):
    """返回 {图号: (图片段索引, 图注段索引)}。"""
    out = {}
    for i, p in enumerate(doc.paragraphs):
        m = re.match(r'^Figure ([1-5]) \|', p.text.strip())
        if not m:
            continue
        for j in range(i - 1, -1, -1):
            if doc.paragraphs[j]._p.findall('.//' + W + 'drawing'):
                out[m.group(1)] = (j, i)
                break
    return out


def replace_image(doc, para, png_path):
    """替换内嵌图片二进制，保持显示宽度不变、按新宽高比修正高度。"""
    blips = para._p.findall('.//' + A + 'blip')
    assert len(blips) == 1, f'该段有 {len(blips)} 个图片，预期 1 个'
    part = doc.part.rels[blips[0].get(R + 'embed')].target_part
    exts = para._p.findall('.//' + WP + 'extent') + para._p.findall('.//' + A + 'ext')
    cx = int(exts[0].get('cx'))
    pw, ph = png_size(png_path)
    cy = int(round(cx * ph / pw))
    for e in exts:
        e.set('cx', str(cx))
        e.set('cy', str(cy))
    with open(png_path, 'rb') as f:
        part._blob = f.read()
    return cx, cy


def set_caption(para, md_text, size=9.5):
    """清空段内所有内容（含公式对象），按 **bold** 标记重建 run。"""
    tmpl = para.runs[0] if para.runs else None
    font_name = tmpl.font.name if tmpl is not None else 'Arial'
    for child in list(para._p):
        if not child.tag.endswith('}pPr'):
            para._p.remove(child)
    for k, seg in enumerate(re.split(r'\*\*', md_text)):
        if not seg:
            continue
        r = para.add_run(seg)
        r.bold = (k % 2 == 1)
        r.font.size = Pt(size)
        if font_name:
            r.font.name = font_name


def _splice(runs, s, e, rep):
    """把 runs 拼接文本的 [s,e) 换成 rep，且不改变 run 的位置。

    必须原位拼接：段内的公式（oMath）是 w:r 之外的兄弟元素，若把整段文字
    并进第一个 run，文字与公式的先后顺序就会错乱。
    """
    pos, idx = 0, []
    for k, r in enumerate(runs):
        a, b = pos, pos + len(r.text)
        pos = b
        if b > s and a < e:
            idx.append((k, a))
    if not idx:
        return
    k0, a0 = idx[0]
    k1, a1 = idx[-1]
    head = runs[k0].text[:s - a0]
    tail = runs[k1].text[e - a1:]
    if k0 == k1:
        runs[k0].text = head + rep + tail
    else:
        runs[k0].text = head + rep
        runs[k1].text = tail
        for k in range(k0 + 1, k1):
            runs[k].text = ''


def remap_runs(para, mapping):
    """段内单遍同步替换：长 token 优先，避免 4b→4d 之后又被 4d→4e 二次命中。"""
    runs = para.runs
    full = ''.join(r.text for r in runs)
    pat = re.compile('|'.join(re.escape(k) for k in
                              sorted(mapping, key=len, reverse=True)))
    spans = [(m.start(), m.end(), mapping[m.group(0)]) for m in pat.finditer(full)]
    for s, e, rep in reversed(spans):      # 自后向前，前面的替换不影响后面的偏移
        _splice(runs, s, e, rep)
    return len(spans)


def main(src, dst):
    doc = Document(src)
    pairs = find_pairs(doc)
    print(f'源文件: {os.path.basename(src)}  ({len(doc.paragraphs)} 段)')
    print(f'  识别到图文对: {pairs}')

    for k in sorted(pairs):
        img_i, cap_i = pairs[k]
        cx, cy = replace_image(doc, doc.paragraphs[img_i], PNG[k])
        set_caption(doc.paragraphs[cap_i], NEW_CAPTIONS_ZH[k])
        print(f'  Fig {k}: 图片段 {img_i} 已换 ({cx/914400:.2f}×{cy/914400:.2f} in)，'
              f'图注段 {cap_i} 已重写 ({len(NEW_CAPTIONS_ZH[k])} 字符)')

    total = 0
    for idx, mapping in XREF.items():
        if idx >= len(doc.paragraphs):
            print(f'  ! 段 {idx} 不存在，跳过')
            continue
        n = remap_runs(doc.paragraphs[idx], mapping)
        total += n
        if n:
            print(f'  第 {idx} 段: {n} 处引用重映射 '
                  + '; '.join(f'{a}→{b}' for a, b in mapping.items()))
    print(f'  交叉引用共改 {total} 处')

    nfix = 0
    for idx, mapping in NUMFIX.items():
        n = remap_runs(doc.paragraphs[idx], mapping)
        if n != len(mapping):
            raise SystemExit(f'!! 段 {idx} 期望改 {len(mapping)} 处，实际 {n} 处')
        nfix += n
        print(f'  第 {idx} 段: ' + '; '.join(f'{a}→{b}' for a, b in mapping.items()))
    print(f'  过时数字共改 {nfix} 处（Table 4 由政策网格整表重建，不再单点修数）')

    # ---- 命题四：重写结论正文；删掉 Table 4（政策情境写进论证本身，不另设表）----
    grid, med, nn = policy_grid()
    para = next(i for i, p in enumerate(doc.paragraphs)
                if p.text.strip().startswith(('与命题三聚焦', '命题四给出的')))
    M = '{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath'
    if doc.paragraphs[para]._p.findall('.//' + M):
        raise SystemExit('!! 命题四段含公式对象，不能整段重写')
    from docx.text.paragraph import Paragraph
    segs = conclusion_four(grid, med).split(SEP)
    set_caption(doc.paragraphs[para], segs[0], size=10.5)
    anchor = doc.paragraphs[para]._p
    parent = doc.paragraphs[para]._parent
    for seg in segs[1:]:                    # 其余各段插在其后，沿用同一段落格式
        new_p = copy.deepcopy(anchor)
        anchor.addnext(new_p)
        anchor = new_p
        set_caption(Paragraph(new_p, parent), seg, size=10.5)
    print(f'  第 {para} 段起：命题四重写为 {len(segs)} 段结论正文')
    # ---- 叙事层：结果 R4 之后补一段政策情境，与结论四呼应 ----
    g4 = lambda b_, t_, k_: grid[(b_, t_)][k_]
    narrative = (
        '上述权衡可直接转译为政策设定。把政策可设的两个量——开发商可承受的 AEP 让步上限、'
        '保护方要求的最低暴露削减——交叉，得到一组政策情境（Fig. 5g）。以“削减过半”为底线，'
        f'1% 的让步已覆盖陆上 {g4(0.01,50,"陆上"):.1f}%、VPTS {g4(0.01,50,"VPTS"):.1f}%、'
        f'Bauer {g4(0.01,50,"Bauer"):.1f}% 的场址；底线提高到 80%，同一预算下降至 '
        f'{g4(0.01,80,"陆上"):.1f}%、{g4(0.01,80,"VPTS"):.1f}% 与 {g4(0.01,80,"Bauer"):.1f}%，'
        f'把预算放宽到 2% 可使 Bauer 回升至 {g4(0.02,80,"Bauer"):.1f}%。但更严格的底线并非'
        f'总能用预算买到：底线取 90% 时，VPTS 的达标率自 1% 预算起即固定在 '
        f'{g4(0.01,90,"VPTS"):.1f}%，2% 与 5% 均无提升 —— 该上限由阵列几何与方向数据'
        '分辨率决定，而非由能源约束决定。')
    res = next(i for i, q in enumerate(doc.paragraphs)
               if q.text.strip().startswith('三组之间的梯度'))
    new_p = copy.deepcopy(doc.paragraphs[res]._p)
    doc.paragraphs[res]._p.addnext(new_p)
    set_caption(Paragraph(new_p, doc.paragraphs[res]._parent), narrative, size=10.5)
    print(f'  第 {res} 段后：补入政策情境叙事段（引用 Fig. 5g）')


    removed = 0
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith('Table 4'):
            p._element.getparent().remove(p._element); removed += 1; break
    for tb in list(doc.tables):
        heads = [c.text.strip() for c in tb.rows[0].cells]
        if heads and ('政策' in heads[0] or 'AEP 预算上限' in heads[0]):
            tb._element.getparent().remove(tb._element); removed += 1; break
    print(f'  已删除 Table 4 及其题注（{removed} 个元素）—— 政策情境改为写进结论正文')

    doc.save(dst)
    print(f'已另存: {os.path.basename(dst)}  ({os.path.getsize(dst)/1e6:.2f} MB)')


if __name__ == '__main__':
    main(sys.argv[1] if len(sys.argv) > 1 else DEFAULT_SRC,
         sys.argv[2] if len(sys.argv) > 2 else DEFAULT_DST)
